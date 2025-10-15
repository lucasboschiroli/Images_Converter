import sys
import os
from pathlib import Path
from PIL import Image
import subprocess

try:
    from docx import Document
    from docx2pdf import convert as docx_to_pdf_convert
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from PyPDF2 import PdfReader, PdfWriter
    from pdf2docx import Converter as PDFConverter
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import openpyxl
    import xlrd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

def convert_image(input_file, output_format):
    try:
        img = Image.open(input_file)
        
        if output_format.lower() in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
            background = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
            img = background
        
        output_file = Path(input_file).stem + f'.{output_format.lower()}'
        img.save(output_file, quality=95)
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error converting image: {e}")
        return False

def convert_video(input_file, output_format):
    try:
        output_file = Path(input_file).stem + f'.{output_format.lower()}'
        
        presets = {
            'mp4': ['-c:v', 'libx264', '-c:a', 'aac', '-strict', 'experimental'],
            'webm': ['-c:v', 'libvpx-vp9', '-c:a', 'libopus'],
            'avi': ['-c:v', 'mpeg4', '-c:a', 'mp3'],
            'mkv': ['-c:v', 'libx264', '-c:a', 'aac'],
            'mov': ['-c:v', 'libx264', '-c:a', 'aac'],
            'gif': ['-vf', 'fps=10,scale=480:-1:flags=lanczos', '-c:v', 'gif'],
        }
        
        codec_args = presets.get(output_format.lower(), ['-c:v', 'libx264', '-c:a', 'aac'])
        
        cmd = ['ffmpeg', '-i', input_file, *codec_args, '-y', output_file]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            print(f"✓ Converted: {output_file}")
            return True
        else:
            print(f"✗ Error: {result.stderr}")
            return False
    except FileNotFoundError:
        print("✗ Error: ffmpeg not installed. Install with: winget install ffmpeg")
        return False
    except Exception as e:
        print(f"✗ Error converting video: {e}")
        return False

def convert_audio(input_file, output_format):
    try:
        output_file = Path(input_file).stem + f'.{output_format.lower()}'
        
        codecs = {
            'mp3': ['-c:a', 'libmp3lame', '-b:a', '192k'],
            'wav': ['-c:a', 'pcm_s16le'],
            'ogg': ['-c:a', 'libvorbis', '-q:a', '5'],
            'flac': ['-c:a', 'flac'],
            'm4a': ['-c:a', 'aac', '-b:a', '192k'],
            'aac': ['-c:a', 'aac', '-b:a', '192k'],
        }
        
        codec_args = codecs.get(output_format.lower(), ['-c:a', 'libmp3lame'])
        
        cmd = ['ffmpeg', '-i', input_file, *codec_args, '-y', output_file]
        
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode == 0:
            print(f"✓ Converted: {output_file}")
            return True
        else:
            print(f"✗ Error: {result.stderr}")
            return False
    except FileNotFoundError:
        print("✗ Error: ffmpeg not installed. Install with: winget install ffmpeg")
        return False
    except Exception as e:
        print(f"✗ Error converting audio: {e}")
        return False

def convert_document(input_file, output_format):
    input_ext = Path(input_file).suffix.lower()
    output_format = output_format.lower()
    
    if output_format == 'pdf':
        if input_ext in ['.docx', '.doc']:
            return docx_to_pdf(input_file)
        elif input_ext == '.txt':
            return txt_to_pdf(input_file)
    elif output_format == 'docx':
        if input_ext == '.pdf':
            return pdf_to_docx(input_file)
        elif input_ext == '.txt':
            return txt_to_docx(input_file)
    elif output_format == 'txt':
        if input_ext == '.pdf':
            return pdf_to_txt(input_file)
        elif input_ext in ['.docx', '.doc']:
            return docx_to_txt(input_file)
    
    print(f"✗ Conversion from {input_ext} to {output_format} not supported")
    return False

def docx_to_pdf(input_file):
    if not DOCX_AVAILABLE:
        print("✗ Error: Install with: pip install python-docx docx2pdf")
        return False
    try:
        output_file = Path(input_file).stem + '.pdf'
        docx_to_pdf_convert(input_file, output_file)
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def pdf_to_docx(input_file):
    if not PDF_AVAILABLE:
        print("✗ Error: Install with: pip install PyPDF2 pdf2docx")
        return False
    try:
        output_file = Path(input_file).stem + '.docx'
        cv = PDFConverter(input_file)
        cv.convert(output_file)
        cv.close()
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def txt_to_pdf(input_file):
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        output_file = Path(input_file).stem + '.pdf'
        c = canvas.Canvas(output_file, pagesize=letter)
        
        with open(input_file, 'r', encoding='utf-8') as f:
            text = f.read()
        
        y = 750
        for line in text.split('\n'):
            if y < 50:
                c.showPage()
                y = 750
            c.drawString(50, y, line[:80])
            y -= 15
        
        c.save()
        print(f"✓ Converted: {output_file}")
        return True
    except ImportError:
        print("✗ Error: Install with: pip install reportlab")
        return False
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def txt_to_docx(input_file):
    if not DOCX_AVAILABLE:
        print("✗ Error: Install with: pip install python-docx")
        return False
    try:
        output_file = Path(input_file).stem + '.docx'
        doc = Document()
        
        with open(input_file, 'r', encoding='utf-8') as f:
            for line in f:
                doc.add_paragraph(line.strip())
        
        doc.save(output_file)
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def pdf_to_txt(input_file):
    if not PDF_AVAILABLE:
        print("✗ Error: Install with: pip install PyPDF2")
        return False
    try:
        output_file = Path(input_file).stem + '.txt'
        reader = PdfReader(input_file)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            for page in reader.pages:
                f.write(page.extract_text())
        
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def docx_to_txt(input_file):
    if not DOCX_AVAILABLE:
        print("✗ Error: Install with: pip install python-docx")
        return False
    try:
        output_file = Path(input_file).stem + '.txt'
        doc = Document(input_file)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            for para in doc.paragraphs:
                f.write(para.text + '\n')
        
        print(f"✓ Converted: {output_file}")
        return True
    except Exception as e:
        print(f"✗ Error: {e}")
        return False

def convert_excel(input_file, output_format):
    if not EXCEL_AVAILABLE:
        print("✗ Error: Install with: pip install openpyxl xlrd")
        return False
    
    input_ext = Path(input_file).suffix.lower()
    output_file = Path(input_file).stem + f'.{output_format.lower()}'
    
    try:
        if input_ext == '.xls' and output_format == 'xlsx':
            workbook = xlrd.open_workbook(input_file)
            new_wb = openpyxl.Workbook()
            new_wb.remove(new_wb.active)
            
            for sheet_name in workbook.sheet_names():
                sheet = workbook.sheet_by_name(sheet_name)
                new_sheet = new_wb.create_sheet(sheet_name)
                
                for row in range(sheet.nrows):
                    for col in range(sheet.ncols):
                        new_sheet.cell(row + 1, col + 1, sheet.cell_value(row, col))
            
            new_wb.save(output_file)
            print(f"✓ Converted: {output_file}")
            return True
        else:
            print(f"✗ Excel conversion from {input_ext} to {output_format} not supported")
            return False
    except Exception as e:
        print(f"✗ Error: {e}")
        return False
    
def detect_file_type(filename):
    ext = Path(filename).suffix.lower()
    
    image_formats = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp', '.ico'}
    video_formats = {'.mp4', '.avi', '.mkv', '.mov', '.wmv', '.flv', '.webm', '.m4v'}
    audio_formats = {'.mp3', '.wav', '.ogg', '.flac', '.m4a', '.aac', '.wma'}
    
    if ext in image_formats:
        return 'image'
    elif ext in video_formats:
        return 'video'
    elif ext in audio_formats:
        return 'audio'
    else:
        return None

def batch_convert(input_dir, output_format):
    files = list(Path(input_dir).glob('*'))
    converted = 0
    
    for file in files:
        if file.is_file():
            file_type = detect_file_type(str(file))
            if file_type:
                print(f"\nConverting: {file.name}")
                if convert_file(str(file), output_format):
                    converted += 1
    
    print(f"\n✓ Converted {converted}/{len(files)} files")

def convert_file(input_file, output_format):
    if not os.path.exists(input_file):
        print(f"✗ Error: File '{input_file}' not found")
        return False
    
    file_type = detect_file_type(input_file)
    
    if file_type == 'image':
        return convert_image(input_file, output_format)
    elif file_type == 'video':
        return convert_video(input_file, output_format)
    elif file_type == 'audio':
        return convert_audio(input_file, output_format)
    elif file_type == 'document':
        return convert_document(input_file, output_format)
    elif file_type == 'excel':
        return convert_excel(input_file, output_format)
    else:
        print(f"✗ Unsupported file type: {Path(input_file).suffix}")
        return False

def main():
    if len(sys.argv) < 3:
        print("Universal File Converter")
        print("\nUsage:")
        print("  Single file:  python converter.py <input_file> <output_format>")
        print("  Batch mode:   python converter.py --batch <directory> <output_format>")
        print("\nSupported Formats:")
        print("  Images:    jpg, png, gif, bmp, tiff, webp, ico")
        print("  Videos:    mp4, avi, mkv, mov, webm, gif")
        print("  Audio:     mp3, wav, ogg, flac, m4a, aac")
        print("  Documents: pdf, docx, txt")
        print("  Excel:     xls, xlsx")
        print("\nExamples:")
        print("  python converter.py image.png jpg")
        print("  python converter.py video.mp4 webm")
        print("  python converter.py document.docx pdf")
        print("  python converter.py file.pdf txt")
        print("  python converter.py --batch ./files pdf")
        sys.exit(1)
    
    if sys.argv[1] == '--batch':
        if len(sys.argv) < 4:
            print("✗ Error: Please provide directory and output format")
            sys.exit(1)
        batch_convert(sys.argv[2], sys.argv[3])
    else:
        input_file = sys.argv[1]
        output_format = sys.argv[2]
        convert_file(input_file, output_format)

if __name__ == '__main__':
    main()